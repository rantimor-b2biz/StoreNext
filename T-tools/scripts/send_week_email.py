#!/usr/bin/env python3
"""
StoreNext — (Re)send the weekly content email for a given week.
Reads O-output/W[NN]/final/ (final-post.md, article.md, article.docx, *-visual.png)
and emails everything to Ran. Used for corrections or missed deliveries.

Usage: python T-tools/scripts/send_week_email.py W27 [--note "..."]
Requires: GMAIL_APP_PASS, GMAIL_EMAIL, RECIPIENT_MAIL. Regenerates article.docx
from article.md if python-docx is available (so corrections to the .md carry over).
"""

import argparse
import os
import re
import smtplib
import sys
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]


def markdown_to_docx(md: str, out_path: Path) -> bool:
    try:
        from docx import Document
    except ImportError:
        return False
    doc = Document()

    def add_runs(par, text):
        for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            run = par.add_run(chunk)
            run.bold = i % 2 == 1

    for block in md.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("### "):
            doc.add_heading(block[4:].strip(), level=3)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=1)
        elif all(l.strip().startswith(("- ", "* ")) for l in block.splitlines()):
            for l in block.splitlines():
                add_runs(doc.add_paragraph(style="List Bullet"), l.strip()[2:])
        else:
            add_runs(doc.add_paragraph(), block.replace("\n", " "))
    doc.save(out_path)
    return True


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("week", help="Week folder, e.g. W27")
    parser.add_argument("--note", default="", help="Optional note shown at the top of the email")
    args = parser.parse_args()

    final_dir = ROOT / "O-output" / args.week / "final"
    post_file = final_dir / "final-post.md"
    if not post_file.exists():
        print(f"{post_file} not found", file=sys.stderr)
        return 1

    # refresh the docx from the (possibly corrected) markdown
    article_md = final_dir / "article.md"
    if article_md.exists():
        markdown_to_docx(article_md.read_text(encoding="utf-8"), final_dir / "article.docx")

    sender = os.environ.get("GMAIL_EMAIL", "rantimor@gmail.com")
    recipient = os.environ.get("RECIPIENT_MAIL", sender)
    password = os.environ["GMAIL_APP_PASS"]

    sep = "=" * 50
    parts = ["שלום רן, ", ""]
    if args.note:
        parts += [f"*** {args.note} ***", ""]
    parts += [
        f"התוכן של {args.week} מצורף (מאמר .docx/.md + ויזואל PNG).",
        "הפוסט המלא למטה.",
        "",
        sep,
        "",
        post_file.read_text(encoding="utf-8"),
    ]

    msg = MIMEMultipart()
    msg["Subject"] = f"StoreNext {args.week} — תוכן מעודכן להעלאה"
    msg["From"] = sender
    msg["To"] = recipient
    msg.attach(MIMEText("\n".join(parts), "plain", "utf-8"))

    for f in sorted(final_dir.iterdir()):
        if f.name == "final-post.md" or not f.is_file():
            continue
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read_bytes())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{f.name}"')
        msg.attach(part)

    with smtplib.SMTP("smtp.gmail.com", 587) as s:
        s.starttls()
        s.login(sender, password)
        s.sendmail(sender, [recipient], msg.as_string())
    print(f"Email sent: {args.week} -> {recipient}")
    return 0


if __name__ == "__main__":
    sys.exit(main())

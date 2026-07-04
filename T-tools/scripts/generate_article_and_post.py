#!/usr/bin/env python3
"""
StoreNext — Automated Article + LinkedIn Post Generator
========================================================
Replaces the old weekly topic-proposals flow (T-tools/generate-draft.py).
Same mechanism as the FRACTIONAL_CMO automation, adapted for StoreNext:
everything is DELIVERED BY EMAIL for manual upload (no auto-publish).

Pipeline (mirrors the agency workflow):
  Stage 1  RESEARCHER   (Sonnet + web search) — trending enterprise topic + sourced facts
  Stage 2  COPYWRITER   (Sonnet)              — full 1,200-1,800 word article (English,
           formal 8/10, data-driven) for the StoreNext website
  Stage 3  GATEKEEPER   (Opus)                — reviews/revises the article, writes the
           LinkedIn post (approved StoreNext formula) + visual-data fields
  Stage 4  ARTIST       (T-tools/generate-visual.py) — brand-exact PNG card (PIL renderer)
  Stage 5  DELIVER      — saves to O-output/W[NN]/ (final/ + process/) and emails
           Ran the post + article (.md + .docx) + visual PNG for manual upload

Usage:
  python T-tools/scripts/generate_article_and_post.py [--topic "..."] [--date YYYY-MM-DD]

Requires: ANTHROPIC_API_KEY, GMAIL_APP_PASS, GMAIL_EMAIL, RECIPIENT_MAIL.
pip install anthropic pillow cairosvg lxml python-docx
"""

import argparse
import datetime
import json
import os
import re
import smtplib
import subprocess
import sys
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

import anthropic

ROOT = Path(__file__).resolve().parents[2]  # StoreNext/

RESEARCH_MODEL = "claude-sonnet-5"    # trend research (needs web search)
DRAFT_MODEL = "claude-sonnet-5"       # article draft (per agency model rules)
GATEKEEPER_MODEL = "claude-opus-4-8"  # final review (per agency model rules)

MAX_TOPIC_HISTORY = 60
RECENT_TOPICS_SHOWN = 15

client = anthropic.Anthropic()


# ---------------------------------------------------------------- helpers

def read_file(rel_path: str) -> str:
    p = ROOT / rel_path
    return p.read_text(encoding="utf-8") if p.exists() else ""


def load_topic_history() -> list:
    p = ROOT / "B-brain" / "topic-history.json"
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return []
    return []


def save_topic_history(history: list) -> None:
    (ROOT / "B-brain" / "topic-history.json").write_text(
        json.dumps(history[-MAX_TOPIC_HISTORY:], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def extract_json(text: str) -> dict:
    fence = re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, re.DOTALL)
    if fence:
        return json.loads(fence.group(1))
    start = text.find("{")
    if start == -1:
        raise ValueError(f"No JSON object found in response:\n{text[:500]}")
    depth, in_str, esc = 0, False, False
    for i in range(start, len(text)):
        c = text[i]
        if in_str:
            if esc:
                esc = False
            elif c == "\\":
                esc = True
            elif c == '"':
                in_str = False
        elif c == '"':
            in_str = True
        elif c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return json.loads(text[start : i + 1])
    raise ValueError("Unbalanced JSON in response")


def text_of(response) -> str:
    return "\n".join(b.text for b in response.content if b.type == "text")


def call_model(model: str, system: str, user_prompt: str, max_tokens: int = 8000):
    with client.messages.stream(
        model=model, max_tokens=max_tokens, system=system,
        messages=[{"role": "user", "content": user_prompt}],
    ) as stream:
        return stream.get_final_message()


def run_with_web_search(system: str, user_prompt: str, max_continuations: int = 5):
    messages = [{"role": "user", "content": user_prompt}]
    tools = [{"type": "web_search_20260209", "name": "web_search", "max_uses": 8}]
    for _ in range(max_continuations):
        response = client.messages.create(
            model=RESEARCH_MODEL, max_tokens=8000, system=system,
            tools=tools, messages=messages,
        )
        if response.stop_reason != "pause_turn":
            return response
        messages.append({"role": "assistant", "content": response.content})
    return response


def iso_week(d: datetime.date) -> str:
    return f"W{d.isocalendar()[1]:02d}"


# ---------------------------------------------------------------- stage 1: research

def stage1_research(topic_override: str | None) -> dict:
    session_brief = read_file("C-core/session-brief.md")
    capabilities = read_file("C-core/product-capabilities.md")
    history = load_topic_history()
    recent = [h.get("topic", "") for h in history[-RECENT_TOPICS_SHOWN:]]

    system = f"""You are the Researcher for StoreNext (enterprise B2B: supplier management +
financial operations for Israeli enterprises).

CLIENT BRIEF:
{session_brief}

APPROVED PRODUCT CAPABILITIES (product_focus MUST match this file's attribution rules):
{capabilities}

CONTENT PILLARS:
1. Supply chain resilience & supplier management (Supplier Portal — "Supply Chain Automations")
2. Financial operations & treasury (Meteor — AI-native "Financial Operations"; direct competitor: NILUS)
3. Procurement & the CFO agenda — cost pressure, efficiency, risk, P2P
4. AI in enterprise finance & procurement — adoption, ROI, governance

YOUR JOB: find what enterprise CFOs / Procurement Directors / CIOs are actually talking
about THIS WEEK, and pick ONE topic that (a) is genuinely trending or newly relevant,
(b) fits one pillar, and (c) supports both a full website article (1,200-1,800 words)
and a LinkedIn post. Use web search to verify — collect 4-6 concrete facts, stats, or
developments WITH named sources (Gartner, Hackett, McKinsey, Deloitte, industry press).
Israeli enterprise context is a plus when natural (never political).

HARD RULES: never frame the two products as one "unified platform"; never use
"40 years" as positioning; never invent statistics.

AVOID repeating these recently covered topics:
{json.dumps(recent, ensure_ascii=False, indent=2)}
"""

    if topic_override:
        task = (
            f'The topic has been chosen manually: "{topic_override}".\n'
            "Research current context, data points, and discussion around it (use web search), "
            "then produce the research brief."
        )
    else:
        task = (
            "Research what is trending right now for enterprise finance and procurement leaders "
            "(recent news, analyst reports, data). Consider 3 candidate topics, then pick the "
            "strongest one."
        )

    task += """

Return your final answer as a JSON object in a ```json fenced block:
{
  "topic": "short topic title",
  "pillar": "which content pillar",
  "product_focus": "Supplier Portal" or "Meteor" or "none (pure thought leadership)",
  "angle": "StoreNext's specific data-driven angle for CFOs/procurement, 2-3 sentences",
  "why_now": "why this is timely THIS week",
  "key_facts": [{"fact": "specific stat or development", "source": "publication name + url"}],
  "article_outline": ["working section heading 1", "...", "4-6 sections"],
  "candidates_considered": ["topic a", "topic b", "topic c"]
}"""

    print("Stage 1: researching trends...", flush=True)
    response = run_with_web_search(system, task)
    brief = extract_json(text_of(response))
    print(f"  -> topic: {brief.get('topic')}", flush=True)
    return brief


# ---------------------------------------------------------------- stage 2: article

def stage2_article(brief: dict) -> dict:
    voice_dna = read_file("C-core/voice-dna.md")
    recent_titles = [h.get("topic", "") for h in load_topic_history()[-RECENT_TOPICS_SHOWN:]]

    system = f"""You are the Copywriter for StoreNext. You write website articles for an
enterprise B2B audience (CFOs, Procurement Directors, CIOs of Israeli enterprises).

VOICE DNA (follow precisely):
{voice_dna}

NON-NEGOTIABLE VOICE RULES:
- English. Professional, formal tone (8/10). Short sentences (10-14 words).
- Data-backed claims with named in-line sources ("According to Gartner, ...").
- Enterprise vocabulary: procurement, vendor, supply chain, ROI, P2P, reconciliation.
- NEVER: "Revolutionary/Transformative/Disruptive", hype, em dashes (use periods),
  exclamation marks, casual language.
- NEVER "unified platform" (products are separate implementations). NEVER "40 years".
- Meteor approved numbers if relevant: 20K+ transactions / 1K+ ERP integrations /
  400+ clients / 150+ global banks.
- Product mentions are capabilities in context, not sales pitches. Solution-focused close.

TITLE RULE: do not reuse naming formulas from recent titles:
{json.dumps(recent_titles, ensure_ascii=False, indent=2)}

ARTICLE STANDARDS:
- 1,200-1,800 words. 4-6 sections with clear headings. One core idea, fully developed.
- Structure: name the problem CFOs feel -> why it exists (structural, not vendor-bashing)
  -> what the data says -> what leading organizations do differently -> practical
  takeaways -> soft solution-focused close."""

    task = f"""Write this week's article from the research brief:

{json.dumps(brief, ensure_ascii=False, indent=2)}

Return JSON in a ```json fenced block:
{{
  "title": "article title",
  "meta_description": "SEO meta description, under 160 chars",
  "category": "e.g. 'Procurement & Supply Chain' / 'Financial Operations'",
  "markdown": "the FULL article in markdown: # Title, then intro paragraphs, ## section headings, **bold** for emphasis, - bullets where useful. No em dashes anywhere."
}}"""

    print("Stage 2: writing article...", flush=True)
    response = call_model(DRAFT_MODEL, system, task, max_tokens=16000)
    article = extract_json(text_of(response))
    wc = len(article.get("markdown", "").split())
    print(f"  -> {article.get('title')} ({wc} words)", flush=True)
    return article


# ---------------------------------------------------------------- stage 3: gatekeeper + post + visual data

def stage3_gatekeeper(brief: dict, article: dict, week: str) -> dict:
    voice_dna = read_file("C-core/voice-dna.md")
    session_brief = read_file("C-core/session-brief.md")
    capabilities = read_file("C-core/product-capabilities.md")

    system = f"""You are the Gatekeeper for StoreNext. Nothing ships without your approval.
StoreNext sells to enterprise CFOs. One hyped sentence destroys credibility.

CLIENT BRIEF:
{session_brief}

FULL VOICE DNA (the standard you enforce):
{voice_dna}

APPROVED PRODUCT CAPABILITIES (the ONLY product claims allowed — if the article
attributes a capability not listed here, or to the wrong product, FIX IT; when in
doubt describe the capability generically without naming a product):
{capabilities}

YOU HAVE THREE JOBS:

A) Review the article. Fix any failure yourself and return the full corrected markdown.
   1. English, formal 8/10, short sentences (10-14 words), data-backed with named sources
   2. Zero: em dashes, exclamation marks, "Revolutionary/Transformative/Disruptive",
      hype, casual language, "unified platform", "40 years"
   3. Serves CFOs/Procurement/CIOs. One core idea. Product tie-in is a capability, not a pitch.

B) Write the LinkedIn post (the approved StoreNext formula):
   - Contrast or tension hook in first 2 lines (first line under 140 chars). Professional,
     never dramatic or exploitative.
   - One idea. Short lines. White space. 1-2 specific stats with named sources.
   - Explain WHY the problem exists (scattered data, manual processes) before any solution.
   - Product appears as "This is exactly where [capability] becomes critical", capabilities
     as outcomes bullets. Never a sales pitch.
   - Close with an insight (not a pitch), then an open question to CFOs/procurement leaders.
   - 5-7 hashtags at the end. NO links in the body. NO em dashes. NO exclamation marks. English.

C) Produce visual-data fields for the brand PNG renderer. Pick visual_type by this logic:
   key metric (%, $, time) -> "stat_card" (default; data-forward performs best with CFOs);
   sequential process/before-after -> "process_flow"; pure insight -> "quote_card"."""

    task = f"""Research brief:
{json.dumps(brief, ensure_ascii=False, indent=2)}

Article to review:
{json.dumps(article, ensure_ascii=False, indent=2)}

Return JSON in a ```json fenced block:
{{
  "verdict": "APPROVED" or "REVISED",
  "review_notes": ["note per check that needed attention, or 'clean pass'"],
  "article": {{ "title": "...", "meta_description": "...", "category": "...", "markdown": "full final article" }},
  "post_text": "the complete LinkedIn post exactly as it should be published, hashtags included",
  "first_comment": "first comment: link placeholder to the article on storenext site + one value-add line",
  "visual": {{
    "post_id": "{week}",
    "visual_type": "stat_card|process_flow|quote_card",
    "category": "short category label",
    "topic": "short visual headline (max 8 words)",
    "key_metric": "the dominant number, e.g. '15-20%'",
    "hook": "one-line hook for the visual (max 18 words)",
    "key_stat": "one-line supporting fact with numbers",
    "source": "named source(s)",
    "visual_direction": "one sentence describing the visual intent"
  }}
}}"""

    print("Stage 3: gatekeeper review + post + visual data...", flush=True)
    response = call_model(GATEKEEPER_MODEL, system, task, max_tokens=24000)
    result = extract_json(text_of(response))
    print(f"  -> verdict: {result.get('verdict')}", flush=True)
    return result


# ---------------------------------------------------------------- stage 4: visual PNG

def stage4_visual(visual: dict, process_dir: Path) -> Path | None:
    vd_path = process_dir / "visual-data.json"
    vd_path.write_text(
        json.dumps({"posts": [visual]}, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print("Stage 4: rendering brand visual (generate-visual.py)...", flush=True)
    try:
        subprocess.run(
            [sys.executable, str(ROOT / "T-tools" / "generate-visual.py"), str(vd_path)],
            check=True, cwd=ROOT,
        )
        png = process_dir.parent / "final" / f"{visual.get('post_id', 'post')}-visual.png"
        if png.exists():
            print(f"  -> {png.name}", flush=True)
            return png
    except subprocess.CalledProcessError as e:
        print(f"  visual rendering failed: {e}", flush=True)
    return None


# ---------------------------------------------------------------- docx conversion

def markdown_to_docx(md: str, out_path: Path) -> bool:
    try:
        from docx import Document
    except ImportError:
        return False

    doc = Document()

    def add_runs(par, text):
        for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            run = par.add_run(chunk)
            run.bold = i % 2 == 1

    for block in md.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("### "):
            doc.add_heading(block[4:].strip(), level=3)
        elif block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=1)
        elif all(l.strip().startswith(("- ", "* ")) for l in block.splitlines()):
            for l in block.splitlines():
                add_runs(doc.add_paragraph(style="List Bullet"), l.strip()[2:])
        else:
            add_runs(doc.add_paragraph(), block.replace("\n", " "))
    doc.save(out_path)
    return True


# ---------------------------------------------------------------- stage 5: email

def send_email(week: str, brief: dict, review: dict, attachments: list[Path]) -> None:
    sender = os.environ.get("GMAIL_EMAIL", "rantimor@gmail.com")
    recipient = os.environ.get("RECIPIENT_MAIL", sender)
    password = os.environ.get("GMAIL_APP_PASS", "")
    if not password:
        print("GMAIL_APP_PASS not set. Skipping email.", flush=True)
        return

    article = review["article"]
    sep = "=" * 50
    body = "\n".join([
        "שלום רן,",
        "",
        f"התוכן השבועי של StoreNext ({week}) מוכן להעלאה ידנית:",
        "",
        "1. מאמר לאתר — מצורף (.docx להעלאה נוחה + .md)",
        "2. פוסט לינקדאין — בגוף המייל למטה (העתק-הדבק)",
        "3. ויזואל ממותג — PNG מצורף",
        "",
        "צ'קליסט: העלה את המאמר לאתר -> פרסם את הפוסט עם הויזואל -> הוסף את",
        "לינק המאמר בתגובה הראשונה.",
        "",
        sep,
        f"ARTICLE: {article.get('title', '')}",
        f"Meta description: {article.get('meta_description', '')}",
        sep,
        "",
        "LINKEDIN POST:",
        "",
        review.get("post_text", ""),
        "",
        sep,
        "",
        "FIRST COMMENT (add the live article URL):",
        "",
        review.get("first_comment", ""),
        "",
        sep,
        f"Gatekeeper: {review.get('verdict')} | Pillar: {brief.get('pillar')} | Product: {brief.get('product_focus')}",
    ])

    msg = MIMEMultipart()
    msg["Subject"] = f"StoreNext {week} — מאמר + פוסט + ויזואל מוכנים להעלאה"
    msg["From"] = sender
    msg["To"] = recipient
    msg.attach(MIMEText(body, "plain", "utf-8"))

    for f in attachments:
        if not f or not f.exists():
            continue
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read_bytes())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{f.name}"')
        msg.attach(part)

    with smtplib.SMTP("smtp.gmail.com", 587) as s:
        s.starttls()
        s.login(sender, password)
        s.sendmail(sender, [recipient], msg.as_string())
    print(f"Stage 5: email sent to {recipient}", flush=True)


# ---------------------------------------------------------------- main

def main() -> int:
    parser = argparse.ArgumentParser(description="Generate StoreNext article + LinkedIn post")
    parser.add_argument("--topic", help="Manual topic override")
    parser.add_argument("--date", help="Date YYYY-MM-DD (default: today)")
    args = parser.parse_args()

    date = datetime.date.fromisoformat(args.date) if args.date else datetime.date.today()
    week = iso_week(date)

    final_dir = ROOT / "O-output" / week / "final"
    process_dir = ROOT / "O-output" / week / "process"
    final_dir.mkdir(parents=True, exist_ok=True)
    process_dir.mkdir(parents=True, exist_ok=True)

    brief = stage1_research(args.topic)
    article_draft = stage2_article(brief)
    review = stage3_gatekeeper(brief, article_draft, week)
    article = review.get("article") or article_draft
    review["article"] = article
    visual = review.get("visual", {})
    visual.setdefault("post_id", week)

    png = stage4_visual(visual, process_dir)

    # --- files (final-post.md format keeps the Firebase mediaPlan sync working)
    publish_str = date.isoformat()
    (final_dir / "final-post.md").write_text(
        "\n".join([
            f"# StoreNext LinkedIn {week}",
            "",
            "**Status:** approved",
            f"**Category:** {visual.get('category', article.get('category', ''))}",
            f"**Publish:** {publish_str}",
            "**Language:** English",
            "",
            "---",
            "",
            review.get("post_text", ""),
            "",
            "---",
            "",
            f"**First comment:** {review.get('first_comment', '')}",
            "",
            f"**Article:** {article.get('title', '')} (attached to the email; upload to the website)",
        ]),
        encoding="utf-8",
    )
    article_md = final_dir / "article.md"
    article_md.write_text(article.get("markdown", ""), encoding="utf-8")
    article_docx = final_dir / "article.docx"
    if not markdown_to_docx(article.get("markdown", ""), article_docx):
        article_docx = None
    (process_dir / "research-brief.md").write_text(
        f"# Research Brief — {week} ({publish_str})\n\n```json\n{json.dumps(brief, ensure_ascii=False, indent=2)}\n```\n",
        encoding="utf-8",
    )
    (process_dir / "gatekeeper-review.md").write_text(
        f"# Gatekeeper Review — {week}\n\n**Verdict:** {review.get('verdict')}\n\n"
        + "\n".join(f"- {n}" for n in review.get("review_notes", [])) + "\n",
        encoding="utf-8",
    )
    (process_dir / "content-process-log.md").write_text(
        "\n".join([
            f"# Content Process Log — {week}",
            "",
            f"- Date: {publish_str}",
            f"- Topic: {brief.get('topic')}",
            f"- Pillar: {brief.get('pillar')} | Product focus: {brief.get('product_focus')}",
            "- Pipeline: automated (Researcher -> Copywriter -> Gatekeeper -> Artist -> Email)",
            f"- Gatekeeper verdict: {review.get('verdict')}",
            "- Delivery: email to Ran for manual upload (article + post + visual)",
        ]),
        encoding="utf-8",
    )

    history = load_topic_history()
    history.append({
        "date": publish_str,
        "week": week,
        "topic": brief.get("topic"),
        "pillar": brief.get("pillar"),
        "product_focus": brief.get("product_focus"),
    })
    save_topic_history(history)

    attachments = [article_md, article_docx, png]
    send_email(week, brief, review, [a for a in attachments if a])

    gh_output = os.environ.get("GITHUB_OUTPUT")
    if gh_output:
        with open(gh_output, "a", encoding="utf-8") as f:
            f.write(f"week={week}\n")
            f.write(f"topic={brief.get('topic')}\n")
            f.write(f"verdict={review.get('verdict')}\n")

    print(f"\nDone. {week}: {brief.get('topic')}")
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except anthropic.APIStatusError as e:
        print(f"Anthropic API error {e.status_code}: {e.message}", file=sys.stderr)
        sys.exit(1)
